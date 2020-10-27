from Modelo.Enums.TipoRO import Tipo
import datetime
import os.path
import docx
import re


class RegistroDeOcorrenciaWord:

    def __init__(self, arquivo):

        self.instrumento_contratual = None
        self.contrato_sap = None
        self.orgao = None
        self.contratada = None
        self.autorizacao_servico = None
        self.prazo_contratual = None
        self.obj_contrato = None
        self.inicio = None
        self.termino = None
        self.local_execucao = None
        self.numero = None
        self.tipo = None
        self.data = None
        self.corpo_fiscalizacao = None
        self.corpo_contratada = None
        self.assinatura_fiscalizacao = None
        self.assinatura_contratada = None
        self.app_word = None
        self.endereco_arquivo = arquivo

    @property
    def instrumento_contratual(self):
        return self.__instrumento_contratual

    @instrumento_contratual.setter
    def instrumento_contratual(self, value):
        if value:
            try:
                valor = value.split('\n')
                value = int(valor[1].replace('.', ''))
                self.__instrumento_contratual = value
            except Exception as e:
                raise Exception("Instrumento Contratual inválido\n"
                                "Erro: {}.".format(e))

    @property
    def contrato_sap(self):
        return self.__contrato_sap

    @contrato_sap.setter
    def contrato_sap(self, value):
        if value:
            try:
                valor = value.split('\n')
                value = int(valor[3].replace('.', ''))
                self.__contrato_sap = value
            except Exception as e:
                raise Exception("Contrato SAP inválido\n"
                                "Erro: {}.".format(e))

    @property
    def orgao(self):
        return self.__orgao

    @orgao.setter
    def orgao(self, value):
        if value:
            try:
                valor = value.split('\n')
                value = valor[1].upper()
                self.__orgao = value
            except Exception as e:
                raise Exception("Órgão inválido\n"
                                "Erro: {}.".format(e))

    @property
    def contratada(self):
        return self.__contratada

    @contratada.setter
    def contratada(self, value):
        if value:
            try:
                valor = value.split('\n')
                value = valor[1].upper()
                self.__contratada = value
            except Exception as e:
                raise Exception("Contratada inválido\n"
                                "Erro: {}.".format(e))

    @property
    def autorizacao_servico(self):
        return self.__autorizacao_servico

    @autorizacao_servico.setter
    def autorizacao_servico(self, value):
        if value:
            try:
                valor = value.split('\n')
                value = datetime.datetime.strptime(valor[1], '%d/%m/%Y')
                self.__autorizacao_servico = value.date()
            except Exception as e:
                raise Exception("Autorização Serviço inválido\n"
                                "Erro: {}.".format(e))

    @property
    def prazo_contratual(self):
        return self.__prazo_contratual

    @prazo_contratual.setter
    def prazo_contratual(self, value):
        if value:
            try:
                valor = value.split('\n')
                value = int(valor[1].replace(' DIAS',''))
                self.__prazo_contratual = value
            except Exception as e:
                raise Exception("Prazo Contratual inválido\n"
                                "Erro: {}.".format(e))

    @property
    def obj_contrato(self):
        return self.__obj_contrato

    @obj_contrato.setter
    def obj_contrato(self, value):
        if value:
            try:
                valor = value.split('\n')
                value = valor[1].upper()
                self.__obj_contrato = value
            except Exception as e:
                raise Exception("Objeto do Contrato inválido\n"
                                "Erro: {}.".format(e))

    @property
    def inicio(self):
        return self.__inicio

    @inicio.setter
    def inicio(self, value):
        if value:
            try:
                valor = value.split('\n')
                value = datetime.datetime.strptime(valor[1], '%d.%m.%Y')
                self.__inicio = value.date()
            except Exception as e:
                raise Exception("Inicio inválido\n"
                                "Erro: {}.".format(e))

    @property
    def termino(self):
        return self.__termino

    @termino.setter
    def termino(self, value):
        if value:
            try:
                valor = value.split('\n')
                value = datetime.datetime.strptime(valor[1], '%d.%m.%Y')
                self.__termino = value.date()
            except Exception as e:
                raise Exception("Término inválido\n"
                                "Erro: {}.".format(e))

    @property
    def local_execucao(self):
        return self.__local_execucao

    @local_execucao.setter
    def local_execucao(self, value):
        if value:
            try:
                valor = value.split('\n')
                value = valor[1].upper()
                self.__local_execucao = value
            except Exception as e:
                raise Exception("Local Execução inválido\n"
                                "Erro: {}.".format(e))

    @property
    def numero(self):
        return self.__numero

    @numero.setter
    def numero(self, value):
        if value:
            value = int(value.replace('NÚMERO\n', ''))
            self.__numero = value

    @property
    def tipo(self):
        return self.__tipo

    @tipo.setter
    def tipo(self, value):
        if value:
            value = re.sub(r"\s+", "", value.upper())
            if value[-1] == 'X':
                self.__tipo = Tipo.Outros
            elif value[6] == 'X':
                self.__tipo = Tipo.Diario
            else:
                self.__tipo = Tipo.Semanal

    @property
    def data(self):
        return self.__data

    @data.setter
    def data(self, value):
        if value:
            value = datetime.datetime.strptime(value.replace('DATA\n', ''), '%d/%m/%Y')
            self.__data = value.date()

    @property
    def corpo_fiscalizacao(self):
        return self.__corpo_fiscalizacao

    @corpo_fiscalizacao.setter
    def corpo_fiscalizacao(self, value):
        self.__corpo_fiscalizacao = value

    @property
    def corpo_contratada(self):
        return self.__corpo_contratada

    @corpo_contratada.setter
    def corpo_contratada(self, value):
        self.__corpo_contratada = value

    @property
    def assinatura_fiscalizacao(self):
        return self.__assinatura_fiscalizacao

    @assinatura_fiscalizacao.setter
    def assinatura_fiscalizacao(self, value):
        self.__assinatura_fiscalizacao = value

    @property
    def assinatura_contratada(self):
        return self.__assinatura_contratada

    @assinatura_contratada.setter
    def assinatura_contratada(self, value):
        self.__assinatura_contratada = value

    @property
    def endereco_arquivo(self):
        return self.__endereco_arquivo

    @endereco_arquivo.setter
    def endereco_arquivo(self, value):
        if value:
            if os.path.exists(value):
                self.__endereco_arquivo = value
            else:
                raise Exception("O Arquivo não foi encontrado!\n"
                                "Arquivo: {}.".format(value))

    @property
    def app_word(self):
        return self.__app_word

    @app_word.setter
    def app_word(self, value):
        self.__app_word = value

    def carrega_arquivo(self):
        try:
            self.app_word = docx.Document(self.endereco_arquivo)
            self.instrumento_contratual = self.app_word.tables[0].cell(1, 3).text
            self.contrato_sap = self.app_word.tables[0].cell(1, 3).text
            self.orgao = self.app_word.tables[0].cell(2, 0).text
            self.contratada = self.app_word.tables[0].cell(2, 3).text
            self.autorizacao_servico = self.app_word.tables[0].cell(2, 4).text
            self.prazo_contratual = self.app_word.tables[0].cell(3, 1).text
            self.obj_contrato = self.app_word.tables[0].cell(3, 3).text
            self.inicio = self.app_word.tables[0].cell(4, 0).text
            self.termino = self.app_word.tables[0].cell(4, 1).text
            self.local_execucao = self.app_word.tables[0].cell(4, 2).text
            self.numero = self.app_word.tables[0].cell(0, 3).text
            self.data = self.app_word.tables[0].cell(0, 4).text
            self.tipo = self.app_word.tables[0].cell(1, 0).text
            self.corpo_fiscalizacao = self.app_word.tables[0].cell(6, 0).text
            self.corpo_contratada = self.app_word.tables[0].cell(6, 3).text
            self.assinatura_fiscalizacao = self.app_word.tables[0].cell(8, 0).text
            self.assinatura_contratada = self.app_word.tables[0].cell(8, 3).text
        except Exception as e:
            raise Exception(e)

    def __str__(self):
        obj = 'Numero:................... {0} \n' + \
              'Data:..................... {1} \n' + \
              'Tipo:..................... {2} \n' + \
              'Corpo Fiscalização:....... {3} \n' + \
              'Corpo Contratada:......... {4} \n' + \
              'Assinatura Fiscalização:.. {5} \n' + \
              'Assinatura Contratada:.... {6} \n' + \
              'Instrumento Contratual:................................... {7} \n' + \
              'Contrato SAP:............................................. {8} \n' + \
              'Órgão:.................................................... {9} \n' + \
              'Contratada:............................................... {10} \n' + \
              'Autorização de Serviço (AS):.............................. {11} \n' + \
              'Prazo Contratual:......................................... {12} Dias \n' + \
              'Objeto do Contrato / Instalação / Natureza dos Serviços:.. {13} \n' + \
              'Início:................................................... {14} \n' + \
              'Término:.................................................. {15} \n' + \
              'Local de Execução dos Serviços:........................... {16} \n'
        return obj.format(
            self.numero, self.data, self.tipo.name, self.corpo_fiscalizacao,
            self.corpo_contratada, self.assinatura_fiscalizacao, self.assinatura_contratada,
            self.instrumento_contratual, self.contrato_sap, self.orgao, self.contratada,
            self.autorizacao_servico, self.prazo_contratual, self.obj_contrato, self.inicio,
            self.termino, self.local_execucao)
